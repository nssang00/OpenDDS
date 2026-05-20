from docx import Document


def generate_docx(title, content, output="output.docx"):
    doc = Document()

    doc.add_heading(title, level=1)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save(output)

    print(f"Saved: {output}")
